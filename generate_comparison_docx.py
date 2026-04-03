"""
Generate Comparison_Table.docx from the latest simulation results.
Reads CSV/JSON results from all three paper implementations.
"""
import json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def style_header_row(table, row_idx=0, color="2F5496"):
    """Style a header row with dark background and white text."""
    for cell in table.rows[row_idx].cells:
        set_cell_shading(cell, color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

def add_data_row(table, row_idx):
    """Center-align data cells."""
    for cell in table.rows[row_idx].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
    style_header_row(table)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
        add_data_row(table, r_idx + 1)
        # Alternate row shading
        if r_idx % 2 == 1:
            for cell in table.rows[r_idx + 1].cells:
                set_cell_shading(cell, "EAF0F9")

    return table


# ---- Load all results ----
with open("Base-Paper-Implementation/all_results.json", "r") as f:
    base = json.load(f)
with open("Forward-Secure-Lattice-Based-Encryption-For-Internet-Of-Things/all_results.json", "r") as f:
    proposed = json.load(f)
with open("Research-paper-2-Implementation-for-comparison-/all_results.json", "r") as f:
    oo_iribe = json.load(f)

# ---- Create Document ----
doc = Document()

# Title
title = doc.add_heading('Performance Comparison: Three-Way Analysis', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ---- Papers Compared ----
doc.add_heading('Research Papers Compared', level=1)
add_table(doc,
    ["#", "Scheme", "Paper", "Description"],
    [
        ["1", "Proposed", "fs-IBE + Novel Trust Model", "fs-IBE with Dilithium-3 trust, epoch-bound queries, FTAR"],
        ["2", "Base Paper", "fs-IBE (Pure)", "fs-IBE only — no trust model, no signatures"],
        ["3", "OO-IRIBE-EnDKER", "Scientific Reports 2025", "Online/Offline IBE with revocation, Number List, cloud decryption"],
    ]
)

doc.add_paragraph()

# ======================================================================
# TABLE 1: Parameter-Level Performance
# ======================================================================
doc.add_heading('Table 1: Parameter-Level Performance Comparison', level=1)
doc.add_paragraph('All schemes tested with: num_data=5, num_queries=10, tree_depth=3')

param_names = ["PARA.512", "PARA.768", "PARA.1024"]
metrics_keys = [
    ("data_encryption_time_s", "Data Encryption Time (s)"),
    ("query_encryption_time_s", "Query Encryption Time (s)"),
    ("data_decryption_time_s", "Data Decryption Time (s)"),
    ("query_execution_latency_s", "Query Execution Latency (s)"),
    ("query_throughput_per_s", "Query Throughput (queries/s)"),
    ("overall_model_throughput_per_s", "Overall Model Throughput (ops/s)"),
    ("overall_model_latency_s", "Overall Model Latency (s)"),
    ("false_trust_acceptance_rate", "False Trust Acceptance Rate"),
]

for p_idx, pname in enumerate(param_names):
    p_meta = {
        "PARA.512": "(n=512, q=3329, NIST Level 1, 143-bit security)",
        "PARA.768": "(n=768, q=3329, NIST Level 3, 207-bit security)",
        "PARA.1024": "(n=1024, q=3329, NIST Level 5, 272-bit security)",
    }
    doc.add_heading(f'{pname} {p_meta[pname]}', level=2)

    prop_m = proposed["parameter_metrics"][p_idx]
    base_m = base["parameter_metrics"][p_idx]
    oo_m = oo_iribe["parameter_metrics"][p_idx]

    rows = []
    for key, label in metrics_keys:
        pv = prop_m.get(key, "N/A")
        bv = base_m.get(key, "N/A")
        ov = oo_m.get(key, "N/A")

        if key == "false_trust_acceptance_rate":
            pv_str = f"{float(pv):.2%}" if pv not in ("N/A", None) else "N/A"
            bv_str = "N/A" if bv == "N/A" else f"{float(bv):.2%}"
            ov_str = f"{float(ov):.2%}" if ov not in ("N/A", None) else "N/A"
        elif isinstance(pv, float):
            pv_str = f"{pv:.4f}"
            bv_str = f"{float(bv):.4f}" if bv != "N/A" else "N/A"
            ov_str = f"{float(ov):.4f}"
        else:
            pv_str, bv_str, ov_str = str(pv), str(bv), str(ov)

        rows.append([label, pv_str, bv_str, ov_str])

    add_table(doc,
        ["Metric", "Proposed (fs-IBE + Trust)", "Base Paper (fs-IBE Pure)", "OO-IRIBE-EnDKER"],
        rows
    )
    doc.add_paragraph()

# ======================================================================
# TABLE 2: Device-Count Metrics
# ======================================================================
doc.add_heading('Table 2: Performance Metrics by Device Count', level=1)
doc.add_paragraph('All schemes tested at n=512, with device counts: 20, 40, 60, 80, 100')

device_counts = [20, 40, 60, 80, 100]

# 2A: Auth Latency
doc.add_heading('2A: Authentication Latency — seconds', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['auth_latency_s']:.4f}",
        f"{base_d['auth_latency_s']:.4f}",
        f"{oo_d['auth_latency_s']:.4f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# 2B: Computation Cost
doc.add_heading('2B: Computation Cost — milliseconds', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['computation_cost_ms']:,.2f}",
        f"{base_d['computation_cost_ms']:,.2f}",
        f"{oo_d['computation_cost_ms']:,.2f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# 2C: Throughput
doc.add_heading('2C: Throughput — operations per second', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['throughput_ops_s']:.2f}",
        f"{base_d['throughput_ops_s']:.2f}",
        f"{oo_d['throughput_ops_s']:.2f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# 2D: Normalized Throughput
doc.add_heading('2D: Verified Normalized Throughput (per device)', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['normalized_throughput']:.3f}",
        f"{base_d['normalized_throughput']:.3f}",
        f"{oo_d['normalized_throughput']:.3f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# 2E: Storage Overhead
doc.add_heading('2E: Storage Overhead — KB', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['storage_overhead_kb']:,.2f}",
        f"{base_d['storage_overhead_kb']:,.2f}",
        f"{oo_d['storage_overhead_kb']:,.2f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# ======================================================================
# TABLE 3: Batch Processing
# ======================================================================
doc.add_heading('Table 3: Batch Processing Metrics (seconds)', level=1)

# 3A: Batch Formation
doc.add_heading('3A: Batch Formation Time', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['batch_formation_s']:.4f}",
        f"{base_d['batch_formation_s']:.4f}",
        f"{oo_d['batch_formation_s']:.4f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# 3B: Batch Decryption
doc.add_heading('3B: Batch Decryption Time', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['batch_decryption_s']:.4f}",
        f"{base_d['batch_decryption_s']:.4f}",
        f"{oo_d['batch_decryption_s']:.4f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# 3C: Batch Authentication
doc.add_heading('3C: Batch Authentication Time', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['batch_authentication_s']:.4f}",
        f"{base_d['batch_authentication_s']:.6f}",
        f"{oo_d['batch_authentication_s']:.4f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# ======================================================================
# TABLE 4: Token Metrics
# ======================================================================
doc.add_heading('Table 4: Token Metrics (seconds)', level=1)

# 4A: Token Generation
doc.add_heading('4A: Token Generation Time', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['token_generation_s']:.4f}",
        f"{base_d['token_generation_s']:.4f}",
        f"{oo_d['token_generation_s']:.4f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# 4B: Token Encryption
doc.add_heading('4B: Token Encryption Time', level=2)
rows = []
for i, nd in enumerate(device_counts):
    prop_d = proposed["device_metrics"][i]
    base_d = base["device_metrics"][i]
    oo_d = oo_iribe["device_metrics"][i]
    rows.append([str(nd),
        f"{prop_d['token_encryption_s']:.4f}",
        f"{base_d['token_encryption_s']:.4f}",
        f"{oo_d['token_encryption_s']:.4f}"])
add_table(doc, ["No. of Devices", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"], rows)
doc.add_paragraph()

# ======================================================================
# TABLE 5: Security Features
# ======================================================================
doc.add_heading('Table 5: Security Features Comparison', level=1)
add_table(doc,
    ["Security Feature", "Proposed", "Base Paper", "OO-IRIBE-EnDKER"],
    [
        ["Post-Quantum Security", "✓ Lattice-based (LWE)", "✓ Lattice-based (LWE)", "✓ Lattice-based (LWE)"],
        ["Forward Security", "✓ Binary tree epochs", "✓ Binary tree epochs", "✗ Number List based"],
        ["Identity-Based Encryption", "✓ Dual Regev IBE", "✓ Dual Regev IBE", "✓ Custom IBE"],
        ["Trust Verification", "✓ Dilithium-3 signatures", "✗ Not present", "✓ Signature-based"],
        ["FTAR", "0.00%", "N/A", "0.00%"],
        ["User Revocation", "✗ Not implemented", "✗ Not implemented", "✓ Number List (O(1))"],
        ["Online/Offline Encryption", "✗ Single-phase", "✗ Single-phase", "✓ Split enc"],
        ["Cloud-Assisted Decryption", "✗ Not present", "✗ Not present", "✓ Semi-trusted cloud"],
        ["NIST Security Levels", "1, 3, 5", "1, 3, 5", "1, 3, 5"],
    ]
)
doc.add_paragraph()

# ======================================================================
# Conclusion
# ======================================================================
doc.add_heading('Conclusion', level=1)
add_table(doc,
    ["Criterion", "Best Scheme"],
    [
        ["Fastest Encryption", "Base Paper"],
        ["Best Throughput at Scale", "Base Paper / Proposed"],
        ["Lowest Storage", "OO-IRIBE-EnDKER"],
        ["Fastest Decryption", "Proposed / Base Paper"],
        ["Security Features", "Proposed (trust + forward security)"],
        ["Overall Balance", "Proposed (best security-performance tradeoff)"],
    ]
)

p = doc.add_paragraph()
p.add_run('\nThe Proposed scheme ').bold = False
run = p.add_run('offers the best balance between security and performance')
run.bold = True
p.add_run(' — it includes trust verification (0% FTAR), forward security, and maintains consistent throughput at scale, with only moderate overhead compared to the base paper. OO-IRIBE-EnDKER, while providing useful features like online/offline split encryption and O(1) revocation, incurs significantly higher computation costs that limit its scalability.')

# ---- Save ----
doc.save("Comparison_Table.docx")
print("Saved: Comparison_Table.docx")
